import streamlit as st
import login


login.generarLogin()
if 'usuario' in st.session_state:
    st.header('Informe :green[SEO]')


import streamlit as st
from dotenv import load_dotenv
import os
from openai import OpenAI
from langchain_openai.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_pinecone import PineconeVectorStore
from langchain_openai.embeddings import OpenAIEmbeddings
from pinecone import Pinecone
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from docx import Document
from io import BytesIO

# Load environment variables
load_dotenv()

# OpenAI client
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# Pinecone API key
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")

# Initialize the OpenAI chat model
model = ChatOpenAI(openai_api_key=api_key, model="gpt-4o")

# Initialize the output parser
parser = StrOutputParser()

# Define the prompt template
template = """
Eres un experto realizando informes y analizando los datos. Los textos que te pasan son de un seminarios impartidos por ADIPA. Trata de ser lo más espécifico de acuerdo a la información que contiene el texto. No hagas suposiciones.

Context: {context}

Question: {question}
"""

prompt = ChatPromptTemplate.from_template(template)

# Streamlit UI
st.title("Pinecone Chat")

# Initialize Pinecone
database = Pinecone(api_key=PINECONE_API_KEY)

# List available indexes
indexes = database.list_indexes().names()
selected_index = st.selectbox("Seleccionar base de datos vectorial", indexes)

responses = {}

if selected_index:
    pinecone_index = database.Index(selected_index)
    st.write(f"Conectado a la base de datos: {selected_index}")

    # Define the first question input
    question1 = "Write in Spanish. You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."

    # Define the second question input
    question2 = "Write in Spanish. You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."

    if question1:
        # Initialize PineconeVectorStore with embeddings
        embeddings = OpenAIEmbeddings()
        
        # Define documents and index_name
        documents = []  # Replace with your actual documents
        index_name = selected_index

        pinecone = PineconeVectorStore.from_documents(
            documents, embeddings, index_name=index_name
        )

        # Define the retrieval and processing chain for the first question
        chain1 = (
            {"context": pinecone.as_retriever(), "question": RunnablePassthrough()}  # Retrieve context from Pinecone
            | prompt  # Apply the prompt template
            | model  # Use the model to generate a response
            | parser  # Parse the model's response
        )

        # Invoke the chain with the first question
        response1 = chain1.invoke(question1)
        responses["Objetivo general"] = response1
        # Display the response for the first question
        st.text_area("Objetivo general:", response1, height=150, key="response1")

    if question2:
        # Define the retrieval and processing chain for the second question
        chain2 = (
            {"context": pinecone.as_retriever(), "question": RunnablePassthrough()}  # Retrieve context from Pinecone
            | prompt  # Apply the prompt template
            | model  # Use the model to generate a response
            | parser  # Parse the model's response
        )

        # Invoke the chain with the second question
        response2 = chain2.invoke(question2)
        responses["Puntos principales"] = response2
        # Display the response for the second question
        st.text_area("Puntos principales:", response2, height=150, key="response2")

    # Define the third question input
    question3 = "Write in Spanish. You are an expert SEO manager. Extract the six most important keywords from this text. Provide only the keywords."

    if question3:
        # Define the retrieval and processing chain for the third question
        chain3 = (
            {"context": pinecone.as_retriever(), "question": RunnablePassthrough()}  # Retrieve context from Pinecone
            | prompt  # Apply the prompt template
            | model  # Use the model to generate a response
            | parser  # Parse the model's response
        )

        # Invoke the chain with the third question
        response3 = chain3.invoke(question3)
        responses["Palabras clave importantes"] = response3
        # Display the response for the third question
        st.text_area("Palabras clave importantes:", response3, height=150, key="response3")

    # Define the fourth question input
    question4 = "Write in Spanish. You are an expert SEO manager. Suggest headings (H1 and H2) for SEO. Provide only the headings."

    if question4:
        # Define the retrieval and processing chain for the fourth question
        chain4 = (
            {"context": pinecone.as_retriever(), "question": RunnablePassthrough()}  # Retrieve context from Pinecone
            | prompt  # Apply the prompt template
            | model  # Use the model to generate a response
            | parser  # Parse the model's response
        )

        # Invoke the chain with the fourth question
        response4 = chain4.invoke(question4)
        responses["Sugerencias de encabezados (H1 y H2)"] = response4
        # Display the response for the fourth question
        st.text_area("Sugerencias de encabezados (H1 y H2):", response4, height=150, key="response4")
    
    # Define the fifth question input
    question5 = "Escribe un blog de entre 800 a 1200 palabras hablando del texto, optimizado para SEO. Asegúrate de incluir palabras clave relevantes, encabezados (H1, H2, H3), y mantener una estructura clara y coherente. El objetivo es crear un contenido atractivo y útil que pueda posicionarse bien en los motores de búsqueda."

    if question5:
        # Define the retrieval and processing chain for the fifth question
        chain5 = (
            {"context": pinecone.as_retriever(), "question": RunnablePassthrough()}  # Retrieve context from Pinecone
            | prompt  # Apply the prompt template
            | model  # Use the model to generate a response
            | parser  # Parse the model's response
        )

        # Invoke the chain with the fifth question
        response5 = chain5.invoke(question5)
        responses["Blog optimizado para SEO"] = response5
        # Display the response for the fifth question
        st.text_area("Blog optimizado para SEO:", response5, height=300, key="response5")

    # Create a Word document with the responses
    doc = Document()
    doc.add_heading('Informe ADIPIO el más Bacan', 0)

    for title, content in responses.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Provide a download button
    st.download_button(
        label="Descargar Informe",
        data=buffer,
        file_name="informe_adipio_chat.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )